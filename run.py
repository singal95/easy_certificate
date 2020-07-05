"""a simple script for automatically generating docx file of certificates"""

# encoding=utf-8
# !/usr/bin/env python
import os
import time
# from math import sqrt

from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import arg_parser, tid_maker


def load_images():
    """return paths of images of pwd"""
    image_files = list(f for f in os.listdir() if f.endswith('.png') or f.endswith('.jpg') or f.endswith('.bmp'))
    if image_files:
        print('load images {}'.format(image_files))
        return image_files
    else:
        return None


def run(rows, collumns, margin, orientation="horizontal"):
    """main excecution func"""
    if rows >= 10 or rows < 0:
        raise ValueError("每行图片太多！请设置较小于等于10的数值")
    if collumns >= 10 or collumns < 0:
        raise ValueError("每列图片太多！请设置小于等于10的数值")
    if margin > 10:
        raise ValueError("页边距太大！请设置小于10的页边距")

    # load images
    image_files = load_images()

    if image_files is None:
        raise IOError("此目录下无图片，请放置图片后重启此程序！")
    image_nums = len(image_files)

    # start to write document in docx format
    document = Document()

    # controls page properties
    section = document.sections[0]
    if orientation == 'horizontal':
        height = 21.0
        width = 29.7
        section.page_height = Cm(height)
        section.page_width = Cm(width)
    else:
        height = 29.7
        width = 21.0
        section.page_height = Cm(height)
        section.page_width = Cm(width)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(margin)
    section.bottom_margin = Cm(margin)
    section.left_margin = Cm(margin)
    section.right_margin = Cm(margin)

    # write tables
    tables = []
    for k in range(image_nums):
        table = document.add_table(rows, collumns)
        table.style = 'Table Grid'
        tables.append(table)
        for i in range(rows):
            for j in range(collumns):
                # set paragraph properties
                para = table.cell(i, j).paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                cell = para.add_run()

                # add image
                cell.add_picture(
                    image_files[k],
                    width=Cm((width - 2 * margin - 0.5) / collumns),
                    height=Cm((height - 2 * margin - 0.5) / rows))

    docx_name = tid_maker() + '.docx'
    document.save(docx_name)


if __name__ == "__main__":
    args = arg_parser()
    run(
        rows=args.rows,
        collumns=args.collumns,
        margin=args.margin,
        orientation=args.orientation
    )
